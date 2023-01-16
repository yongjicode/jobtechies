from flask import Flask
from flask import request
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import openai
import requests
import json

openai.api_key = "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"

class SectionGenerator:
    def __init__(self):
        pass

    def insertHR(self, paragraph):
        p = paragraph._p  # p is the <w:p> XML element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.insert_element_before(pBdr,
            'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
            'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
            'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
            'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
            'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
            'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
            'w:pPrChange'
        )
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)

    def create_section(self, document, section_name):
        # section_header = document.add_heading(level=5)
        section_header = document.add_paragraph()

        runner = section_header.add_run(f'{section_name}')
        runner.bold = True
        runner.size = Pt(13)

        self.insertHR(section_header)


    def add_item(self, document, title, time, detail, bullet_point_array):
        EXACTLY=1.15
        title_para = document.add_paragraph()
        runner = title_para.add_run(title)
        runner.bold = True
        
        title_para.add_run(f'\t{time}')  # tab will trigger tabstop
        sec = document.sections[0]
        # finding end_point for the content 
        margin_end = Inches(
            sec.page_width.inches - (sec.left_margin.inches + sec.right_margin.inches))
        tab_stops = title_para.paragraph_format.tab_stops
        # adding new tab stop, to the end point, and making sure that it's `RIGHT` aligned.
        tab_stops.add_tab_stop(margin_end, WD_TAB_ALIGNMENT.RIGHT)
        title_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        title_para.paragraph_format.space_after = Pt(2)

        b=document.add_paragraph(f"{detail}")
        b.paragraph_format.space_after = Pt(2)
        b.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        for i in bullet_point_array:
            c=document.add_paragraph(i, style='List Bullet')
            c.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            c.paragraph_format.space_after = Pt(2)

    def add_edu_item(self, document, school_name, time, type_of_education, grades,others):
        EXACTLY=1.15
        title_para = document.add_paragraph()
        runner = title_para.add_run(school_name)
        runner.bold = True
        
        title_para.add_run(f'\t{time}')  # tab will trigger tabstop
        sec = document.sections[0]
        # finding end_point for the content 
        margin_end = Inches(
            sec.page_width.inches - (sec.left_margin.inches + sec.right_margin.inches))
        tab_stops = title_para.paragraph_format.tab_stops
        # adding new tab stop, to the end point, and making sure that it's `RIGHT` aligned.
        tab_stops.add_tab_stop(margin_end, WD_TAB_ALIGNMENT.RIGHT)
        title_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        title_para.paragraph_format.space_after = Pt(2)

        b=document.add_paragraph(f"{type_of_education}")
        b.paragraph_format.space_after = Pt(2)
        b.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        
        c=document.add_paragraph(f"Current GPA: {grades}", style='List Bullet')
        c.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        c.paragraph_format.space_after = Pt(2)
        c=document.add_paragraph(others, style='List Bullet')

    def add_work_item(self, document, coy_name, time, role, description):
        EXACTLY=1.15
        coy_name_para = document.add_paragraph()
        runner = coy_name_para.add_run(coy_name)
        runner.bold = True
        
        coy_name_para.add_run(f'\t{time}')  # tab will trigger tabstop
        sec = document.sections[0]
        # finding end_point for the content 
        margin_end = Inches(
            sec.page_width.inches - (sec.left_margin.inches + sec.right_margin.inches))
        tab_stops = coy_name_para.paragraph_format.tab_stops
        # adding new tab stop, to the end point, and making sure that it's `RIGHT` aligned.
        tab_stops.add_tab_stop(margin_end, WD_TAB_ALIGNMENT.RIGHT)
        coy_name_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        coy_name_para.paragraph_format.space_after = Pt(2)

        b=document.add_paragraph(f"{role}")
        b.paragraph_format.space_after = Pt(2)
        b.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        c=document.add_paragraph(description, style='List Bullet')

    def add_volunteer_item(self, document, title, time, description):
        EXACTLY=1.15
        coy_name_para = document.add_paragraph()
        runner = coy_name_para.add_run(title)
        runner.bold = True
        
        coy_name_para.add_run(f'\t{time}')  # tab will trigger tabstop
        sec = document.sections[0]
        # finding end_point for the content 
        margin_end = Inches(
            sec.page_width.inches - (sec.left_margin.inches + sec.right_margin.inches))
        tab_stops = coy_name_para.paragraph_format.tab_stops
        # adding new tab stop, to the end point, and making sure that it's `RIGHT` aligned.
        tab_stops.add_tab_stop(margin_end, WD_TAB_ALIGNMENT.RIGHT)
        coy_name_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        coy_name_para.paragraph_format.space_after = Pt(2)
        c=document.add_paragraph(description, style='List Bullet')


    def create_skills(self, document, language,skills):
        section_header = document.add_heading(level=5)
        runner = section_header.add_run('Skills')
        runner.bold = True
        self.insertHR(section_header)
        para = document.add_paragraph(f"Languages: {language}\n")
        para.add_run(f"Skills: {skills}")

    def create_hobbies(self, document, hobbies):
        section_header = document.add_heading(level=5)
        runner = section_header.add_run('Hobbies')
        runner.bold = True
        self.insertHR(section_header)
        document.add_paragraph(hobbies)

def improveText(text):

    # print(openai.Completion.create(
    #     model="text-davinci-003",
    #     prompt=f"Given the sentence: {text} Make it more convincing for a resume.",
    #     max_tokens=1000,
    #     # temperature=0
    # ).choices[0].text.replace("\n", "").replace('"', ""))

    return openai.Completion.create(
        model="text-davinci-003",
        prompt=f"Given the sentence: {text} Make it more convincing for a resume.",
        max_tokens=1000,
        # temperature=0
    ).choices[0].text.replace("\n", "").replace('"', "")          

app = Flask(__name__)

@app.route("/")
def hello_world():
    return "<p>Hello, World!</p>"

@app.route("/convert", methods=["POST"])
def convert():
    ### Writing The Resume ###

    # resume_dict = {"user_key":{section:{qns:ans}}}

    # print(request.data)

    # resume_dict = {'currentQn': {0: 'skillsInterest', 1: 3}, 'info': {1: 'hi', 2: '97365732', 3: 'nyjyongjie@gmail.com', 4: 'yong-ng@linkedin'}, 'education': {1: 'NTU', 2: 'Business and Computing', 3: '5.0', 4: '100921 - 300525', 5: 'scholarship, statistics and hackathon'}, 'experience': {1: 'ninja van', 2: 'software engineer intern', 3: '123456', 4: 'managed task, blah blah'}, 'cca': {1: 'basketball', 2: 'captain', 3: '123456-1243456', 4: 'played in the interhall game blah blah'}, 'volunteer': {1: 'willing hearts', 2: '098765-09876', 3: 'cut potatoes and vegetable'}, 'skillsInterest': {1: 'english, chinese', 2: 'data science and machine learning', 3: 'software, coding, projects, basketball'}}
    #print(request.data.decode('utf8').replace("'", '"'))
    # print(json.loads(request.data.decode('utf8')))
    #print(request.get_json().data)
    resume_dict = json.loads(request.data.decode('utf8'))
    
    info=resume_dict["info"]
    education=resume_dict["education"]
    education["5"] = improveText(education["5"])
    experience=resume_dict["experience"]
    experience["4"] = improveText(experience["4"])
    cca=resume_dict["cca"]
    cca["4"] = improveText(cca["4"])
    volunteer=resume_dict["volunteer"]
    volunteer["3"] = improveText(volunteer["3"])
    skills_interests=resume_dict["skillsInterest"]

    # data = {"name": "Bryan Lim Kai Wen", "mobile": "+65 12345678", "email": "rando@ntu.edu.sg"}
    document = Document()
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(1.3)
        section.right_margin = Cm(1.3)

    document.add_paragraph('')
    # write resume header info
    title= document.add_paragraph()
    title_format = title.paragraph_format
    title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(f'{info["1"]}\n')
    title_runner=title.add_run(f'Mobile No.: {info["2"]} | Email: {info["3"]} | {info["4"]}.com')
    title_runner.bold = True
    title_runner.size = Pt(10)

    section_generator = SectionGenerator()


    section_generator.create_section(document, "Education")
    section_generator.add_edu_item(document, education["1"], education["4"], education["2"], education["3"], education["5"])

    section_generator.create_section(document, "Experience")
    section_generator.add_work_item(document,experience["1"],experience["3"],experience["2"],experience["4"])  

    section_generator.create_section(document, "CCA")
    section_generator.add_work_item(document,cca["1"],cca["3"],cca["2"],cca["4"])

    section_generator.create_section(document, "Volunteer")
    section_generator.add_volunteer_item(document,volunteer["1"],volunteer["2"],volunteer["3"])

    section_generator.create_skills(document, skills_interests["1"],skills_interests["2"])
    section_generator.create_hobbies(document, skills_interests["3"])

    document.save(f'{resume_dict["info"]["1"]}_Resume.docx')
    
    # r = requests.get(url = "https://api.openai.com/v1/models", headers={'Authorization': "Bearer " + openai.api_key}) # params = PARAMS
    # print("hishfsd: ", r.json())
    # document.save(f'{user_data["info"][1]}_Resume.docx')

    # generate_resume(request.data)
    return "hi"

if __name__ == "__main__":
    app.run()